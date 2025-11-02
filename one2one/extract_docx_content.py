#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
从Word文档中提取一对一门徒训练的内容
"""

from docx import Document
import json

def extract_content_from_docx(docx_path):
    """从Word文档中提取内容"""
    
    doc = Document(docx_path)
    
    content = {
        'preface': [],
        'discipleship_steps': [],
        'lesson1_sections': []
    }
    
    current_section = None
    current_subsection = []
    
    print(f"文档包含 {len(doc.paragraphs)} 个段落")
    print("\n" + "="*80)
    
    for i, para in enumerate(doc.paragraphs):
        text = para.text.strip()
        
        if not text:
            continue
            
        # 打印所有内容以便分析
        print(f"{i}: [{para.style.name}] {text[:80]}")
    
    return content

def main():
    docx_path = "/Users/andyshengruilee/Documents/website/web2Lord/one2one/一对一20251011Andy排版.docx"
    
    print("=" * 80)
    print("开始读取Word文档...")
    print("=" * 80)
    
    content = extract_content_from_docx(docx_path)
    
    # 保存为JSON文件
    output_file = "/Users/andyshengruilee/Documents/website/web2Lord/one2one/docx_content.json"
    with open(output_file, 'w', encoding='utf-8') as f:
        json.dump(content, f, ensure_ascii=False, indent=2)
    
    print("\n" + "=" * 80)
    print(f"提取完成！结果已保存到: {output_file}")
    print("=" * 80)

if __name__ == "__main__":
    main()
