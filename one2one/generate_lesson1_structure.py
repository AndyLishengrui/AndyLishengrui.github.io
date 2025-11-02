#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
根据Word文档重新生成第一课的HTML文件
保留完整内容,在经文处添加填空互动
"""

from docx import Document
import json
import re

def extract_lesson1_sections(docx_path):
    """提取第一课的所有章节内容"""
    
    doc = Document(docx_path)
    
    # 第一课从索引40开始,到120结束
    lesson1_start = 40
    lesson1_end = 120
    
    sections = []
    current_section = None
    
    for i in range(lesson1_start, lesson1_end):
        if i >= len(doc.paragraphs):
            break
            
        para = doc.paragraphs[i]
        text = para.text.strip()
        
        if not text:
            continue
        
        style = para.style.name
        
        # 识别主要章节标题
        if style == 'Heading 2' or any(keyword in text for keyword in [
            '罪使我们与神隔绝',
            '答案：我们罪使我们与神隔绝',
            '解决方案：耶稣基督的牺牲与代罪',
            '耶稣基督十字架的救恩',
            '罪人的回应：凭信心领受神的恩典',
            '个人应用'
        ]):
            if current_section:
                sections.append(current_section)
            current_section = {
                'title': text,
                'content': []
            }
        elif current_section:
            current_section['content'].append({
                'type': 'subtitle' if style in ['Heading 3', 'Heading 4'] else
                        'verse' if style == '经文' else 'text',
                'text': text,
                'style': style
            })
        else:
            # 引言部分
            if not sections:
                sections.append({
                    'title': '引言',
                    'content': [{
                        'type': 'text' if style != '经文' else 'verse',
                        'text': text,
                        'style': style
                    }]
                })
    
    if current_section:
        sections.append(current_section)
    
    return sections

def create_verse_blanks(verse_text):
    """
    为经文创建填空
    智能识别关键词并创建填空
    """
    
    # 定义关键词列表
    keywords = {
        '神': 1, '耶和华': 2, '耶稣': 2, '基督': 2,
        '罪': 1, '罪人': 2, '罪孽': 2, '罪恶': 2,
        '救': 1, '救主': 2, '救赎': 2, '拯救': 2,
        '信': 1, '相信': 2, '信心': 2,
        '永生': 2, '灭亡': 2, '死': 1, '死亡': 2,
        '恩典': 2, '怜悯': 2, '慈爱': 2,
        '义': 1, '公义': 2, '圣洁': 2,
        '血': 1, '十字架': 3
    }
    
    result_text = verse_text
    blanks = []
    
    # 按长度排序,先替换长词
    sorted_keywords = sorted(keywords.items(), key=lambda x: len(x[0]), reverse=True)
    
    for keyword, priority in sorted_keywords:
        if keyword in result_text:
            # 只替换第一次出现
            blank_id = len(blanks)
            blank_marker = f'<blank id="{blank_id}" answer="{keyword}"/>'
            result_text = result_text.replace(keyword, blank_marker, 1)
            blanks.append({'id': blank_id, 'answer': keyword})
            
            if len(blanks) >= 5:  # 每节经文最多5个空
                break
    
    return result_text, blanks

def main():
    docx_path = "/Users/andyshengruilee/Documents/website/web2Lord/one2one/一对一20251011Andy排版.docx"
    
    print("提取第一课内容...")
    sections = extract_lesson1_sections(docx_path)
    
    print(f"提取了 {len(sections)} 个章节")
    
    # 将章节分配到6个小节
    # S1: 引言 + 罪使我们与神隔绝(开头部分)
    # S2: 罪使我们与神隔绝(详细解释)
    # S3: 解决方案
    # S4: 救恩结果  
    # S5: 信心回应
    # S6: 个人应用
    
    section_distribution = {
        'S1': [0, 1],  # 引言 + 第一部分
        'S2': [2],     # 罪的详细解释
        'S3': [3],     # 解决方案
        'S4': [4],     # 救恩结果
        'S5': [5],     # 信心回应
        'S6': [6]      # 个人应用
    }
    
    # 保存为JSON供后续使用
    output_data = {
        'sections': sections,
        'distribution': section_distribution
    }
    
    output_file = "/Users/andyshengruilee/Documents/website/web2Lord/one2one/lesson1_structured.json"
    with open(output_file, 'w', encoding='utf-8') as f:
        json.dump(output_data, f, ensure_ascii=False, indent=2)
    
    print(f"\n结构化数据已保存到: {output_file}")
    
    # 打印章节概览
    print("\n章节概览:")
    print("-"*80)
    for i, section in enumerate(sections):
        content_types = {}
        for item in section['content']:
            content_types[item['type']] = content_types.get(item['type'], 0) + 1
        
        print(f"{i}. {section['title']}")
        print(f"   内容: {dict(content_types)}")

if __name__ == "__main__":
    main()
