#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
对比Word文档内容和HTML文件内容
"""

from bs4 import BeautifulSoup
from docx import Document
import os

def extract_html_content(html_path):
    """从HTML文件提取内容"""
    
    with open(html_path, 'r', encoding='utf-8') as f:
        soup = BeautifulSoup(f.read(), 'html.parser')
    
    # 提取标题
    title = soup.find('h1')
    title_text = title.get_text(strip=True) if title else ''
    
    # 提取描述
    description = soup.find(class_='description')
    desc_text = description.get_text(strip=True) if description else ''
    
    # 提取所有经文
    verses = []
    for verse_section in soup.find_all(class_='verse-section'):
        verse_text_elem = verse_section.find(class_='verse-text')
        if verse_text_elem:
            verse_text = verse_text_elem.get_text(strip=True)
            verses.append(verse_text)
    
    return {
        'title': title_text,
        'description': desc_text,
        'verses': verses,
        'file': os.path.basename(html_path)
    }

def get_docx_section_content(docx_path, section_number):
    """从Word文档提取特定章节内容"""
    
    doc = Document(docx_path)
    
    # 定义章节映射
    section_mapping = {
        1: (40, 48),   # 引言和第一部分
        2: (48, 77),   # 罪使我们与神隔绝
        3: (77, 91),   # 解决方案
        4: (91, 101),  # 我们得救并与神和好
        5: (101, 106), # 凭信心领受
        6: (106, 120)  # 个人应用
    }
    
    if section_number not in section_mapping:
        return []
    
    start, end = section_mapping[section_number]
    content = []
    
    for i in range(start, end):
        if i < len(doc.paragraphs):
            para = doc.paragraphs[i]
            text = para.text.strip()
            if text:
                content.append({
                    'text': text,
                    'style': para.style.name
                })
    
    return content

def compare_section(html_path, docx_path, section_num):
    """对比单个章节"""
    
    print(f"\n{'='*80}")
    print(f"对比文件: {os.path.basename(html_path)}")
    print(f"{'='*80}\n")
    
    html_content = extract_html_content(html_path)
    docx_content = get_docx_section_content(docx_path, section_num)
    
    print(f"HTML标题: {html_content['title']}")
    print(f"HTML描述: {html_content['description'][:100]}..." if len(html_content['description']) > 100 else f"HTML描述: {html_content['description']}")
    print(f"HTML经文数: {len(html_content['verses'])}")
    print(f"Word段落数: {len(docx_content)}")
    
    # 检查经文
    print(f"\n经文对比:")
    print("-"*80)
    
    docx_verses = [c for c in docx_content if c['style'] == '经文']
    
    if len(html_content['verses']) != len(docx_verses):
        print(f"⚠️  经文数量不匹配!")
        print(f"   HTML: {len(html_content['verses'])} 节")
        print(f"   Word: {len(docx_verses)} 节")
    else:
        print(f"✅ 经文数量匹配: {len(html_content['verses'])} 节")
    
    # 显示前3节经文对比
    for i in range(min(3, len(html_content['verses']), len(docx_verses))):
        print(f"\n第{i+1}节经文:")
        print(f"  HTML: {html_content['verses'][i][:80]}...")
        print(f"  Word: {docx_verses[i]['text'][:80]}...")
    
    return html_content, docx_content

def main():
    base_path = "/Users/andyshengruilee/Documents/website/integrated-site/one2one"
    docx_path = "/Users/andyshengruilee/Documents/website/web2Lord/one2one/一对一20251011Andy排版.docx"
    
    html_files = [
        ('C1_S1.html', 1),
        ('C1_S2.html', 2),
        ('C1_S3.html', 3),
        ('C1_S4.html', 4),
        ('C1_S5.html', 5),
        ('C1_S6.html', 6)
    ]
    
    print("="*80)
    print("第一课内容对比报告")
    print("="*80)
    
    for html_file, section_num in html_files:
        html_path = os.path.join(base_path, html_file)
        
        if os.path.exists(html_path):
            try:
                compare_section(html_path, docx_path, section_num)
            except Exception as e:
                print(f"\n❌ 处理 {html_file} 时出错: {e}")
        else:
            print(f"\n⚠️  文件不存在: {html_file}")
    
    print(f"\n{'='*80}")
    print("对比完成!")
    print(f"{'='*80}\n")

if __name__ == "__main__":
    main()
