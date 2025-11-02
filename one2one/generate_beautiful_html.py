#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
改进版HTML生成器 - 生成美观的第一课HTML文件
"""

from docx import Document
import json
import re

def generate_beautiful_html(section_data, section_num, total_sections):
    """生成美观的HTML"""
    
    html = f'''<!DOCTYPE html>
<html lang="zh-CN">
<head>
    <meta charset="UTF-8">
    <meta name="viewport" content="width=device-width, initial-scale=1.0">
    <title>{section_data['title']} - 第{section_num}节 | 一对一门徒训练</title>
    <style>
        * {{
            margin: 0;
            padding: 0;
            box-sizing: border-box;
        }}

        body {{
            font-family: -apple-system, BlinkMacSystemFont, 'Segoe UI', 'PingFang SC', 'Hiragino Sans GB', 'Microsoft YaHei', sans-serif;
            background: linear-gradient(135deg, #667eea 0%, #764ba2 100%);
            min-height: 100vh;
            padding: 20px;
            line-height: 1.8;
        }}

        .container {{
            max-width: 1000px;
            margin: 0 auto;
            background: white;
            border-radius: 15px;
            box-shadow: 0 10px 40px rgba(0,0,0,0.2);
            overflow: hidden;
        }}

        header {{
            background: linear-gradient(135deg, #667eea 0%, #764ba2 100%);
            color: white;
            padding: 40px;
            text-align: center;
        }}

        .course-badge {{
            display: inline-block;
            background: rgba(255,255,255,0.2);
            padding: 8px 20px;
            border-radius: 20px;
            font-size: 0.9em;
            margin-bottom: 15px;
        }}

        h1 {{
            font-size: 2.5em;
            margin-bottom: 10px;
            text-shadow: 2px 2px 4px rgba(0,0,0,0.2);
        }}

        .section-info {{
            font-size: 1em;
            opacity: 0.95;
        }}

        .content {{
            padding: 50px;
        }}

        .section-title {{
            font-size: 2em;
            color: #667eea;
            margin-bottom: 30px;
            padding-bottom: 15px;
            border-bottom: 3px solid #667eea;
            text-align: center;
        }}

        .content-block {{
            margin-bottom: 40px;
        }}

        .text-paragraph {{
            font-size: 1.15em;
            line-height: 2.2;
            color: #2c3e50;
            margin-bottom: 20px;
            text-align: justify;
            padding: 0 20px;
        }}

        .subtitle {{
            font-size: 1.5em;
            font-weight: 600;
            color: #764ba2;
            margin: 40px 0 25px 0;
            padding-left: 20px;
            border-left: 5px solid #764ba2;
        }}

        .verse-container {{
            background: linear-gradient(135deg, #f8f9ff 0%, #fff 100%);
            padding: 30px;
            border-radius: 12px;
            margin: 30px 0;
            border-left: 5px solid #667eea;
            box-shadow: 0 3px 15px rgba(0,0,0,0.08);
        }}

        .verse-text {{
            font-size: 1.2em;
            line-height: 2.3;
            color: #1a202c;
            margin-bottom: 20px;
        }}

        .blank {{
            display: inline-block;
            min-width: 70px;
            padding: 3px 12px;
            border-bottom: 3px solid #667eea;
            margin: 0 5px;
            font-weight: 700;
            color: #667eea;
            cursor: pointer;
            transition: all 0.3s;
            background: linear-gradient(135deg, #e8ebff 0%, #f0f0ff 100%);
            border-radius: 4px 4px 0 0;
        }}

        .blank:hover {{
            background: linear-gradient(135deg, #dce0ff 0%, #e8ebff 100%);
            transform: translateY(-2px);
            box-shadow: 0 3px 10px rgba(102,126,234,0.3);
        }}

        .blank.filled {{
            background: linear-gradient(135deg, #e8f5ff 0%, #e0f2fe 100%);
            border-bottom-color: #0ea5e9;
            color: #0369a1;
        }}

        .blank.correct {{
            background: linear-gradient(135deg, #d1fae5 0%, #bbf7d0 100%);
            border-bottom-color: #10b981;
            color: #065f46;
            animation: correctPulse 0.5s ease-in-out;
        }}

        .blank.wrong {{
            background: linear-gradient(135deg, #fee2e2 0%, #fecaca 100%);
            border-bottom-color: #ef4444;
            color: #991b1b;
            animation: shake 0.5s ease-in-out;
        }}

        @keyframes correctPulse {{
            0%, 100% {{ transform: scale(1); }}
            50% {{ transform: scale(1.05); }}
        }}

        @keyframes shake {{
            0%, 100% {{ transform: translateX(0); }}
            25% {{ transform: translateX(-5px); }}
            75% {{ transform: translateX(5px); }}
        }}

        .verse-reference {{
            text-align: right;
            font-style: italic;
            color: #718096;
            font-size: 1em;
            font-weight: 500;
            margin-top: 15px;
            padding-top: 10px;
            border-top: 1px solid #e2e8f0;
        }}

        .input-section {{
            position: sticky;
            bottom: 0;
            background: white;
            padding: 25px 50px;
            border-top: 3px solid #667eea;
            box-shadow: 0 -5px 30px rgba(0,0,0,0.1);
            z-index: 100;
        }}

        .input-controls {{
            display: flex;
            gap: 15px;
            align-items: center;
            max-width: 1000px;
            margin: 0 auto;
        }}

        #answer-input {{
            flex: 1;
            padding: 15px 25px;
            border: 3px solid #667eea;
            border-radius: 10px;
            font-size: 1.1em;
            transition: all 0.3s;
        }}

        #answer-input:focus {{
            outline: none;
            border-color: #764ba2;
            box-shadow: 0 0 0 4px rgba(102,126,234,0.1);
        }}

        .btn {{
            padding: 15px 30px;
            border: none;
            border-radius: 10px;
            font-size: 1em;
            cursor: pointer;
            font-weight: 600;
            transition: all 0.3s;
            box-shadow: 0 3px 10px rgba(0,0,0,0.1);
        }}

        .btn:hover {{
            transform: translateY(-2px);
            box-shadow: 0 5px 20px rgba(0,0,0,0.15);
        }}

        .btn-check {{
            background: linear-gradient(135deg, #667eea 0%, #764ba2 100%);
            color: white;
        }}

        .btn-reset {{
            background: linear-gradient(135deg, #94a3b8 0%, #64748b 100%);
            color: white;
        }}

        .btn-submit {{
            background: linear-gradient(135deg, #10b981 0%, #059669 100%);
            color: white;
        }}

        .score-display {{
            position: fixed;
            top: 30px;
            right: 30px;
            background: white;
            padding: 20px 35px;
            border-radius: 15px;
            box-shadow: 0 10px 40px rgba(0,0,0,0.2);
            font-weight: 600;
            z-index: 1000;
            border: 3px solid #667eea;
        }}

        .score-label {{
            font-size: 0.95em;
            color: #718096;
            margin-bottom: 5px;
        }}

        .score-number {{
            font-size: 2.5em;
            color: #667eea;
            font-weight: 700;
        }}

        .navigation {{
            display: flex;
            justify-content: space-between;
            padding: 30px 50px;
            background: #f8f9fa;
            border-top: 2px solid #e2e8f0;
        }}

        .nav-btn {{
            padding: 15px 35px;
            background: white;
            color: #667eea;
            text-decoration: none;
            border-radius: 10px;
            border: 3px solid #667eea;
            font-weight: 600;
            transition: all 0.3s;
            box-shadow: 0 3px 10px rgba(0,0,0,0.05);
        }}

        .nav-btn:hover {{
            background: linear-gradient(135deg, #667eea 0%, #764ba2 100%);
            color: white;
            transform: translateY(-2px);
            box-shadow: 0 5px 20px rgba(102,126,234,0.3);
        }}

        .feedback {{
            margin-top: 20px;
            padding: 20px;
            border-radius: 10px;
            text-align: center;
            font-weight: 600;
            font-size: 1.1em;
            display: none;
            animation: slideDown 0.3s ease-out;
        }}

        @keyframes slideDown {{
            from {{
                opacity: 0;
                transform: translateY(-10px);
            }}
            to {{
                opacity: 1;
                transform: translateY(0);
            }}
        }}

        .feedback.show {{
            display: block;
        }}

        .feedback.success {{
            background: linear-gradient(135deg, #d1fae5 0%, #bbf7d0 100%);
            color: #065f46;
            border: 2px solid #10b981;
        }}

        .feedback.error {{
            background: linear-gradient(135deg, #fee2e2 0%, #fecaca 100%);
            color: #991b1b;
            border: 2px solid #ef4444;
        }}

        .feedback.info {{
            background: linear-gradient(135deg, #dbeafe 0%, #bfdbfe 100%);
            color: #1e40af;
            border: 2px solid #3b82f6;
        }}

        @media (max-width: 768px) {{
            .container {{
                border-radius: 0;
            }}
            
            header, .content, .navigation {{
                padding: 25px;
            }}
            
            h1 {{
                font-size: 1.8em;
            }}
            
            .input-section {{
                padding: 20px;
            }}
            
            .input-controls {{
                flex-direction: column;
            }}
            
            .score-display {{
                position: static;
                margin: 20px auto;
                width: fit-content;
            }}

            .text-paragraph {{
                font-size: 1.05em;
                padding: 0 10px;
            }}

            .verse-text {{
                font-size: 1.1em;
            }}
        }}
    </style>
</head>
<body>
    <div class="score-display">
        <div class="score-label">得分</div>
        <div class="score-number" id="score">0</div>
        <div class="score-label" style="font-size: 0.8em;">/ 100</div>
    </div>

    <div class="container">
        <header>
            <div class="course-badge">第一课：新起点 - 得救</div>
            <h1>{section_data['title']}</h1>
            <div class="section-info">第 {section_num} / {total_sections} 节</div>
        </header>

        <div class="content">
'''
    
    # 添加章节标题
    if section_data.get('section_title'):
        html += f'            <h2 class="section-title">{section_data["section_title"]}</h2>\n\n'
    
    blank_id = 0
    blanks_data = []
    
    # 处理内容
    for item in section_data['content']:
        if item['type'] == 'text':
            html += f'            <p class="text-paragraph">{item["text"]}</p>\n\n'
        
        elif item['type'] == 'subtitle':
            html += f'            <h3 class="subtitle">{item["text"]}</h3>\n\n'
        
        elif item['type'] == 'verse':
            # 处理经文填空
            verse_text = item['text']
            reference = item.get('reference', '')
            
            # 简单的填空逻辑:关键词用____替换
            keywords = ['神', '罪', '耶稣', '基督', '救', '信', '义', '恩', '爱', '永生', '独生子', '血', '十字架']
            
            for keyword in keywords:
                if keyword in verse_text and '____' not in verse_text:
                    # 只替换第一次出现
                    verse_text = verse_text.replace(keyword, f'<span class="blank" data-id="{blank_id}">____</span>', 1)
                    blanks_data.append({'id': blank_id, 'answer': keyword})
                    blank_id += 1
            
            html += '            <div class="verse-container">\n'
            html += f'                <div class="verse-text">{verse_text}</div>\n'
            if reference:
                html += f'                <div class="verse-reference">{reference}</div>\n'
            html += '            </div>\n\n'
    
    # 输入控制区域
    html += '''        </div>

        <div class="input-section">
            <div class="input-controls">
                <input type="text" id="answer-input" placeholder="点击填空处,然后输入答案...">
                <button class="btn btn-check" onclick="checkAnswers()">✓ 检查答案</button>
                <button class="btn btn-reset" onclick="resetAll()">↻ 重新开始</button>
                <button class="btn btn-submit" onclick="submitAnswers()">✉ 提交</button>
            </div>
            <div class="feedback" id="feedback"></div>
        </div>

        <div class="navigation">
'''
    
    # 导航按钮
    if section_num > 1:
        html += f'            <a href="one2one_C1_S{section_num-1}.html" class="nav-btn">← 上一节</a>\n'
    else:
        html += '            <a href="index.html" class="nav-btn">← 返回首页</a>\n'
    
    if section_num < total_sections:
        html += f'            <a href="one2one_C1_S{section_num+1}.html" class="nav-btn">下一节 →</a>\n'
    else:
        html += '            <a href="index.html" class="nav-btn">完成 →</a>\n'
    
    html += '''        </div>
    </div>

    <script>
'''
    
    # JavaScript代码
    html += f'        let blanks = {json.dumps(blanks_data, ensure_ascii=False)};\n'
    html += '''        let currentBlankIndex = null;
        let userAnswers = {};

        document.addEventListener('DOMContentLoaded', function() {
            document.querySelectorAll('.blank').forEach(blank => {
                blank.addEventListener('click', function() {
                    document.querySelectorAll('.blank').forEach(b => b.style.boxShadow = 'none');
                    
                    currentBlankIndex = this.dataset.id;
                    this.style.boxShadow = '0 0 0 3px rgba(102,126,234,0.5)';
                    document.getElementById('answer-input').focus();
                });
            });

            document.getElementById('answer-input').addEventListener('keypress', function(e) {
                if (e.key === 'Enter' && currentBlankIndex !== null) {
                    fillBlank();
                }
            });
        });

        function fillBlank() {
            if (currentBlankIndex === null) return;
            
            const input = document.getElementById('answer-input');
            const answer = input.value.trim();
            
            if (!answer) return;
            
            const blank = document.querySelector(`[data-id="${currentBlankIndex}"]`);
            blank.textContent = answer;
            blank.classList.add('filled');
            blank.style.boxShadow = 'none';
            
            userAnswers[currentBlankIndex] = answer;
            
            input.value = '';
            currentBlankIndex = null;
            
            updateScore();
        }

        function checkAnswers() {
            let correct = 0;
            let total = blanks.length;
            let feedback = document.getElementById('feedback');
            
            if (Object.keys(userAnswers).length === 0) {
                feedback.textContent = '请先填写答案！';
                feedback.className = 'feedback show info';
                return;
            }
            
            blanks.forEach(blank => {
                const blankElem = document.querySelector(`[data-id="${blank.id}"]`);
                const userAnswer = userAnswers[blank.id] || '';
                
                blankElem.classList.remove('correct', 'wrong', 'filled');
                
                if (userAnswer === blank.answer) {
                    blankElem.classList.add('correct');
                    correct++;
                } else if (userAnswer) {
                    blankElem.classList.add('wrong');
                    blankElem.title = `正确答案: ${blank.answer}`;
                }
            });
            
            const score = Math.round((correct / total) * 100);
            document.getElementById('score').textContent = score;
            
            if (score === 100) {
                feedback.textContent = `🎉 完美！全部正确！得分：${score}分`;
                feedback.className = 'feedback show success';
            } else if (score >= 80) {
                feedback.textContent = `👍 很好！得分：${score}分 (${correct}/${total}正确)`;
                feedback.className = 'feedback show success';
            } else if (score >= 60) {
                feedback.textContent = `💪 继续加油！得分：${score}分 (${correct}/${total}正确)`;
                feedback.className = 'feedback show info';
            } else {
                feedback.textContent = `📖 需要加强！得分：${score}分 (${correct}/${total}正确)`;
                feedback.className = 'feedback show error';
            }
        }

        function resetAll() {
            userAnswers = {};
            currentBlankIndex = null;
            
            document.querySelectorAll('.blank').forEach(blank => {
                blank.textContent = '____';
                blank.className = 'blank';
                blank.style.boxShadow = 'none';
                blank.removeAttribute('title');
            });
            
            document.getElementById('answer-input').value = '';
            document.getElementById('score').textContent = '0';
            document.getElementById('feedback').className = 'feedback';
        }

        function submitAnswers() {
            checkAnswers();
            
            const score = parseInt(document.getElementById('score').textContent);
            
            if (score === 100) {
                setTimeout(() => {
                    if (confirm('恭喜完成本节！是否继续下一节？')) {
                        document.querySelector('.navigation a:last-child').click();
                    }
                }, 1000);
            }
        }

        function updateScore() {
            let correct = 0;
            let filled = 0;
            
            blanks.forEach(blank => {
                const userAnswer = userAnswers[blank.id];
                if (userAnswer) {
                    filled++;
                    if (userAnswer === blank.answer) {
                        correct++;
                    }
                }
            });
            
            if (filled > 0) {
                const score = Math.round((correct / blanks.length) * 100);
                document.getElementById('score').textContent = score;
            }
        }
    </script>
</body>
</html>
'''
    
    return html

def extract_lesson1_from_docx():
    """从Word文档提取第一课内容并分成6节"""
    doc = Document('/Users/andyshengruilee/Documents/website/web2Lord/one2one/一对一20251011Andy排版.docx')
    
    # 第一课的段落索引范围
    lesson1_ranges = {
        1: {'start': 37, 'end': 45, 'title': '第一节：罪使我们与神隔绝', 'section_title': '罪使我们与神隔绝'},
        2: {'start': 46, 'end': 58, 'title': '第二节:神的解决方法', 'section_title': '神的解决方法'},
        3: {'start': 59, 'end': 68, 'title': '第三节:神爱世人', 'section_title': '神爱世人'},
        4: {'start': 69, 'end': 82, 'title': '第四节:我们的回应', 'section_title': '我们的回应'},
        5: {'start': 83, 'end': 95, 'title': '第五节:如何祷告', 'section_title': '如何祷告'},
        6: {'start': 96, 'end': 109, 'title': '第六节:得救的确据', 'section_title': '得救的确据'}
    }
    
    sections = []
    
    for section_num, info in lesson1_ranges.items():
        content = []
        current_verse = []
        current_verse_ref = ''
        
        for i in range(info['start'], info['end'] + 1):
            if i >= len(doc.paragraphs):
                break
            
            para = doc.paragraphs[i]
            text = para.text.strip()
            
            if not text:
                continue
            
            style = para.style.name
            
            # 处理副标题
            if style == 'Heading 2' or '：' in text[:10]:
                if current_verse:
                    # 保存之前的经文
                    verse_text = ' '.join(current_verse)
                    content.append({
                        'type': 'verse',
                        'text': verse_text,
                        'reference': current_verse_ref
                    })
                    current_verse = []
                    current_verse_ref = ''
                
                content.append({'type': 'subtitle', 'text': text})
            
            # 处理经文
            elif style == '经文' or '(' in text and ('书' in text or '章' in text):
                # 检查是否是经文引用
                if re.search(r'[（(].+[书福利约传音]\s*\d+[:：]\d+', text):
                    # 这是经文引用
                    if current_verse:
                        verse_text = ' '.join(current_verse)
                        content.append({
                            'type': 'verse',
                            'text': verse_text,
                            'reference': text
                        })
                        current_verse = []
                        current_verse_ref = ''
                else:
                    # 这是经文内容
                    current_verse.append(text)
            
            # 处理普通文本
            else:
                if current_verse:
                    # 保存之前的经文
                    verse_text = ' '.join(current_verse)
                    content.append({
                        'type': 'verse',
                        'text': verse_text,
                        'reference': current_verse_ref
                    })
                    current_verse = []
                    current_verse_ref = ''
                
                content.append({'type': 'text', 'text': text})
        
        # 保存最后的经文(如果有)
        if current_verse:
            verse_text = ' '.join(current_verse)
            content.append({
                'type': 'verse',
                'text': verse_text,
                'reference': current_verse_ref
            })
        
        sections.append({
            'title': info['title'],
            'section_title': info['section_title'],
            'content': content
        })
    
    return {'sections': sections}

def main():
    output_dir = "/Users/andyshengruilee/Documents/website/web2Lord/one2one/courses"
    
    print("="*80)
    print("从Word文档提取并生成美观的第一课HTML文件...")
    print("="*80)
    
    # 从Word文档提取内容
    structure = extract_lesson1_from_docx()
    total_sections = len(structure['sections'])
    
    for i, section in enumerate(structure['sections'], 1):
        html_content = generate_beautiful_html(section, i, total_sections)
        
        output_file = f"{output_dir}/one2one_C1_S{i}.html"
        with open(output_file, 'w', encoding='utf-8') as f:
            f.write(html_content)
        
        print(f"✓ 已生成: one2one_C1_S{i}.html - {section['title']}")
    
    print("\n" + "="*80)
    print(f"完成！共生成 {total_sections} 个HTML文件")
    print("="*80)

if __name__ == "__main__":
    main()
