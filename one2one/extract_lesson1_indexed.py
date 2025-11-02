#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
从Word文档中提取第一课得救的完整内容 - 使用段落索引
"""

from docx import Document

def extract_lesson1_by_index(docx_path):
    """根据段落索引提取第一课内容"""
    
    doc = Document(docx_path)
    
    # 根据之前的打印,第一课从索引40开始,到索引119结束(新主人主权之前)
    lesson1_start = 40  # ")新起点 得救"
    lesson1_end = 120   # "新主人 主权"之前
    
    lesson1_text = []
    
    print(f"提取段落 {lesson1_start} 到 {lesson1_end}")
    print("="*80)
    
    for i in range(lesson1_start, lesson1_end):
        if i < len(doc.paragraphs):
            para = doc.paragraphs[i]
            text = para.text.strip()
            
            if text:
                lesson1_text.append({
                    'index': i,
                    'style': para.style.name,
                    'text': text
                })
    
    return lesson1_text

def organize_into_sections(paragraphs):
    """将段落组织成章节"""
    
    sections = []
    current_section = None
    
    for para in paragraphs:
        text = para['text']
        style = para['style']
        
        # 识别主要章节
        if style == 'Heading 2' or '罪使我们与神隔绝' in text or '牺牲与代罪' in text or \
           '我们得救并与神和好' in text or '凭信心领受神的恩典' in text or '个人应用' in text:
            if current_section:
                sections.append(current_section)
            current_section = {
                'title': text,
                'content': []
            }
        elif current_section:
            current_section['content'].append(para)
        else:
            # 引言部分
            if not sections:
                sections.append({
                    'title': '引言',
                    'content': [para]
                })
    
    if current_section:
        sections.append(current_section)
    
    return sections

def main():
    docx_path = "/Users/andyshengruilee/Documents/website/web2Lord/one2one/一对一20251011Andy排版.docx"
    
    print("提取第一课内容...")
    print("="*80)
    
    paragraphs = extract_lesson1_by_index(docx_path)
    
    print(f"\n提取了 {len(paragraphs)} 个段落")
    
    # 组织成章节
    sections = organize_into_sections(paragraphs)
    
    print(f"组织成 {len(sections)} 个章节\n")
    
    # 保存为文本文件
    output_file = "/Users/andyshengruilee/Documents/website/web2Lord/one2one/lesson1_extracted.txt"
    
    with open(output_file, 'w', encoding='utf-8') as f:
        f.write("第一课：新起点 得救\n")
        f.write("="*80 + "\n\n")
        
        for i, section in enumerate(sections, 1):
            f.write(f"\n{'='*80}\n")
            f.write(f"章节 {i}: {section['title']}\n")
            f.write(f"{'='*80}\n\n")
            
            for para in section['content']:
                style = para['style']
                text = para['text']
                
                if style == 'Heading 3' or style == 'Heading 4':
                    f.write(f"\n### {text}\n\n")
                elif style == '经文':
                    f.write(f"【经文】{text}\n\n")
                else:
                    f.write(f"{text}\n\n")
        
        f.write("\n" + "="*80 + "\n")
    
    print(f"完整内容已保存到: {output_file}")
    
    # 打印章节概览
    print("\n章节概览:")
    print("-"*80)
    for i, section in enumerate(sections, 1):
        print(f"{i}. {section['title']} ({len(section['content'])} 个段落)")

if __name__ == "__main__":
    main()
